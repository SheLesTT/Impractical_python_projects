import docx
from docx.shared import RGBColor, Pt

def check_length(real_text, fake_text):
    counter_len = 0
    for paragraph in fake_text .paragraphs:
        if len(paragraph.text) == 0:
            counter_len += 1
    if counter_len < len(real_text.paragraphs):
        print("length of a text is not enought")
    else:
        print("it is ok")

fake_text = docx.Document("fakemessage.docx")
fake_list = []
for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)
    print(paragraph.text)

real_text = docx.Document("realMessage.docx")
real_list = []
for paragraph in real_text.paragraphs:
    if len(paragraph.text) != 0:
        real_list.append(paragraph.text)


check_length(real_text, fake_text)
doc = docx.Document("template.docx")

doc.add_heading("morland Holms", 0)
subtitle = doc.add_heading('Global Consulting Negotiations',1)
subtitle.aligment = 1
doc.add_heading('', 1)
doc.add_paragraph('17 december 2018')
doc.add_paragraph('')

def set_spacing(paragraph):

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

length_real = len(real_list)
count_real = 0

for line in fake_list:
    if count_real < length_real and line == "":
        paragraph = doc.add_paragraph(real_list[count_real])
        paragraph_index = len(doc.paragraphs) - 1

        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255,255,255)

        count_real += 1

    else:
        paragraph = doc.add_paragraph(line)

    set_spacing(paragraph)
doc.save("ciphertext_message.docx")



print("done")

